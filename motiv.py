from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd
import win32com.client
from variables import MY_NAME, MY_EMAIL, DATUM

x = 0


def word_to_pdf(input_path, output_path):
    # Create a new instance of Word application
    word = win32com.client.Dispatch("Word.Application")

    # Open the Word document
    doc = word.Documents.Open(input_path)

    # Save the document as PDF
    # FileFormat 17 represents PDF format
    doc.SaveAs(output_path, FileFormat=17)

    # Close the document and Word application
    doc.Close()
    word.Quit()


def generate_cover_letter(name, address):
    # Create a new Word document
    doc = Document()

    # Adjust document margins
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(40)  # Set left margin to 40 points
        section.right_margin = Pt(40)  # Set right margin to 40 points
        section.top_margin = Pt(40)  # Set top margin to 40 points
        section.bottom_margin = Pt(40)  # Set bottom margin to 40 points

    # Add fixed text for the cover letter
    doc.add_paragraph().add_run(
        f"{MY_NAME}\nAnnaba, Algerien am {DATUM}\n{MY_EMAIL}").bold = True
    address_paragraph = doc.add_paragraph("")
    address_paragraph.add_run(f"{address}\nDeutschland\n\n").bold = True
    address_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add the main title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(
        "Bewerbung um einen Duales Studium Platz :\nGesundheitsmanagement\n")
    title_run.bold = True
    title_run.font.size = Pt(18)  # Set font size to 18 points
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(
        f"\nSehr geehrter {name},")
    doc.add_paragraph("\nhiermit bewerbe ich mich voller Begeisterung um einen Studienplatz im Bereich Gesundheitsmanagement in Ihrem renommierten Unternehmen. Ihre Reputation für exzellente Ausbildung und Fürsorge für Auszubildende und Studenten hat mein Interesse geweckt. Ich bin fest davon überzeugt, dass ich in Ihrem Unternehmen die bestmögliche Ausbildung erhalten kann.")
    doc.add_paragraph("\nIch schätze die inspirierende Arbeitskultur Ihres Unternehmens, die auf Teamwork, Respekt und kontinuierlicher Weiterentwicklung basiert. Ich bin bereit, mich voll und ganz den Herausforderungen und Chancen dieser Ausbildung zu stellen und meine Fähigkeiten unter Ihrer Anleitung weiterzuentwickeln.")
    doc.add_paragraph("\nMeine bisherigen Erfahrungen und Kenntnisse im Bereich der Medizin sowie meine Leidenschaft für Sport, Ernährungswissenschaft und zwischenmenschliche Kommunikation machen mich zu einem geeigneten Kandidaten für diese Stelle.")
    doc.add_paragraph("\nVerantwortung, Teamfähigkeit und Geduld sind Fähigkeiten, die ich erworben habe, während ich mit Menschen gearbeitet habe. Diese Erfahrungen haben in mir die Leidenschaft geweckt, anderen so viel wie möglich zu helfen und sie bestmöglich zu unterstützen.")
    doc.add_paragraph("\nIch bedanke mich herzlich für Ihre Zeit und die Berücksichtigung meiner Bewerbung. Ich freue mich sehr über die Möglichkeit, meine Motivation und Eignung in einem persönlichen Gespräch näher zu erläutern.")
    doc.add_paragraph("\nMit freundlichen Grüßen,\n")

    # Add your name and address at the end of the letter
    doc.add_paragraph(f"{MY_NAME}")

    # Save the document as a Word file
    path = f"Anschreiben.docx"
    doc.save(path)


def CreateMotiv():
    # Read the CSV file into a pandas DataFrame
    df = pd.read_csv('test.csv', encoding="utf-8")

    # Iterate through each row in the DataFrame
    for index, row in df.iterrows():
        # Extract name and address from the current row
        name = row['name']
        address = row['address']
        # Generate the cover letter for the current person

        try:
            generate_cover_letter(name, address)
        except Exception:
            print("An Error happened with :", name, address)
            continue


CreateMotiv()
